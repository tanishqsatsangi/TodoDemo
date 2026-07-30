from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from fpdf import FPDF


ROOT = Path(__file__).resolve().parent
INPUT_MD = ROOT / "STUDY_NOTES_TODO_APP.md"
OUTPUT_DOCX = ROOT / "STUDY_NOTES_TODO_APP.docx"
OUTPUT_PDF = ROOT / "STUDY_NOTES_TODO_APP.pdf"


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def break_long_tokens(text: str, max_len: int = 80) -> str:
    tokens = text.split(" ")
    fixed = []
    for token in tokens:
        if len(token) <= max_len:
            fixed.append(token)
            continue
        chunks = [token[i : i + max_len] for i in range(0, len(token), max_len)]
        fixed.append(" ".join(chunks))
    return " ".join(fixed)


def parse_lines(md_text: str):
    lines = md_text.splitlines()
    parsed = []
    in_code_block = False

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            parsed.append(("code_marker", ""))
            continue

        if in_code_block:
            parsed.append(("code", line))
            continue

        stripped = line.strip()
        if not stripped:
            parsed.append(("blank", ""))
            continue

        if stripped == "---":
            parsed.append(("separator", ""))
            continue

        if stripped.startswith("### "):
            parsed.append(("h3", clean_inline_markdown(stripped[4:])))
        elif stripped.startswith("## "):
            parsed.append(("h2", clean_inline_markdown(stripped[3:])))
        elif stripped.startswith("# "):
            parsed.append(("h1", clean_inline_markdown(stripped[2:])))
        elif stripped.startswith("- "):
            parsed.append(("bullet", clean_inline_markdown(stripped[2:])))
        elif re.match(r"^\d+\.\s+", stripped):
            parsed.append(("number", clean_inline_markdown(stripped)))
        else:
            parsed.append(("text", clean_inline_markdown(stripped)))

    return parsed


def build_docx(parsed, output_path: Path):
    doc = Document()
    doc.add_heading("Todo App Study Notes", level=1)

    for kind, content in parsed:
        if kind == "h1":
            doc.add_heading(content, level=1)
        elif kind == "h2":
            doc.add_heading(content, level=2)
        elif kind == "h3":
            doc.add_heading(content, level=3)
        elif kind == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(content, style="List Number")
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif kind == "separator":
            doc.add_paragraph("-" * 40)
        elif kind == "blank":
            doc.add_paragraph("")
        elif kind == "code_marker":
            # skip marker lines, code lines are already handled
            pass
        else:
            doc.add_paragraph(content)

    doc.save(output_path)


def build_pdf(parsed, output_path: Path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(0, 10, "Todo App Study Notes")
    pdf.ln(2)
    line_width = pdf.w - pdf.l_margin - pdf.r_margin

    def write_line(height: int, text: str):
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(line_width, height, text)

    for kind, content in parsed:
        if kind == "h1":
            pdf.set_font("Helvetica", "B", 16)
            write_line(10, break_long_tokens(content))
        elif kind == "h2":
            pdf.set_font("Helvetica", "B", 14)
            write_line(9, break_long_tokens(content))
        elif kind == "h3":
            pdf.set_font("Helvetica", "B", 12)
            write_line(8, break_long_tokens(content))
        elif kind == "bullet":
            pdf.set_font("Helvetica", "", 11)
            write_line(7, f"- {break_long_tokens(content)}")
        elif kind == "number":
            pdf.set_font("Helvetica", "", 11)
            write_line(7, break_long_tokens(content))
        elif kind == "code":
            pdf.set_font("Courier", "", 10)
            write_line(6, break_long_tokens(content, max_len=70))
        elif kind == "separator":
            pdf.set_font("Helvetica", "", 10)
            write_line(6, "-" * 40)
        elif kind == "blank":
            pdf.ln(2)
        elif kind == "code_marker":
            # skip code fence markers
            pass
        else:
            pdf.set_font("Helvetica", "", 11)
            write_line(7, break_long_tokens(content))

    pdf.output(str(output_path))


def main():
    md_text = INPUT_MD.read_text(encoding="utf-8")
    parsed = parse_lines(md_text)
    docx_path = OUTPUT_DOCX
    pdf_path = OUTPUT_PDF

    try:
        build_docx(parsed, docx_path)
    except PermissionError:
        docx_path = ROOT / "STUDY_NOTES_TODO_APP_v2.docx"
        build_docx(parsed, docx_path)

    try:
        build_pdf(parsed, pdf_path)
    except PermissionError:
        pdf_path = ROOT / "STUDY_NOTES_TODO_APP_v2.pdf"
        build_pdf(parsed, pdf_path)

    print(f"Created: {docx_path}")
    print(f"Created: {pdf_path}")


if __name__ == "__main__":
    main()
